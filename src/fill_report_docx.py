#!/usr/bin/env python3
"""将 report.md 的正文内容填入 report.docx 模板。"""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.text.paragraph import Paragraph

PROJECT_ROOT = Path(__file__).resolve().parent.parent
TEMPLATE_PATH = PROJECT_ROOT / "report_template.docx"
BACKUP_SOURCE = PROJECT_ROOT / "report.docx"
OUTPUT_PATH = PROJECT_ROOT / "report.docx"


def set_paragraph_text(
    paragraph: Paragraph,
    text: str,
    style: str | None = None,
    bold: bool = False,
    align_center: bool = False,
):
    paragraph.text = ""
    run = paragraph.add_run(text)
    run.bold = bold
    if style:
        paragraph.style = style
    if align_center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.space_after = Pt(6)


def insert_paragraph_after(paragraph: Paragraph, doc: Document) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def remove_paragraph(paragraph: Paragraph):
    paragraph._element.getparent().remove(paragraph._element)


def fill_cover_page(doc: Document):
    updates = {
        62: "题 目： 基于 LightGBM 集成学习算法的房价预测模型研究",
        64: "□ 三个以上的基础算法解决经典的仿真问题",
        65: "☑ 扩展算法解决竞赛问题或实际问题",
        66: "□ 提出了创新性的算法思路解决实际问题",
        69: "学        号 \tD202581826",
        70: "姓        名 \t张耀之",
        71: "专        业 \t计算机科学与技术",
        72: "课程指导教师 \t",
        73: "院（系、所） \t计算机科学与技术学院",
        79: "2026年5月1日",
        81: "基于 LightGBM 算法的房价预测模型研究",
        82: "张耀之",
        83: "（计算机科学与技术学院）",
        85: (
            "摘要：房价预测是房地产市场分析中的核心问题，其难点在于影响因素众多、特征间存在复杂的非线性交互关系。"
            "本文基于 Kaggle 经典竞赛数据集 \"House Prices: Advanced Regression Techniques\"，以决策树和随机森林作为基线模型，"
            "以 LightGBM 作为扩展集成学习算法，构建房价回归预测模型。研究首先介绍梯度提升决策树（GBDT）的基本原理，"
            "进而阐述 LightGBM 所采用的直方图加速策略与 leaf-wise 树生长策略。在数据预处理阶段，本文对缺失值进行语义化填充，"
            "构造总面积、总浴室数等领域特征，并对右偏数值特征与目标变量进行对数变换。实验采用十折交叉验证，以 RMSLE 作为评价指标。"
            "结果表明，LightGBM 的平均 RMSLE 为 0.1243，优于决策树（0.2010）和随机森林（0.1398），"
            "验证了扩展算法在复杂房价预测任务中的有效性。"
        ),
        86: "关键词：机器学习；房价预测；LightGBM；集成学习；梯度提升决策树",
    }
    for idx, text in updates.items():
        set_paragraph_text(doc.paragraphs[idx], text)


def clear_body_after_intro(doc: Document) -> Paragraph:
    start_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip().startswith("1 引言"))
    end_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip().startswith("[1]"))
    anchor = doc.paragraphs[start_idx - 1]

    for idx in range(end_idx, start_idx - 1, -1):
        remove_paragraph(doc.paragraphs[idx])

    if len(doc.tables) >= 3:
        old_result_table = doc.tables[2]._tbl
        old_result_table.getparent().remove(old_result_table)

    return anchor


def insert_table_after(doc: Document, anchor: Paragraph, headers, rows, caption: str | None = None) -> Paragraph:
    current = anchor
    if caption:
        current = insert_paragraph_after(current, doc)
        set_paragraph_text(current, caption, bold=True)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for r, row in enumerate(rows, start=1):
        for c, value in enumerate(row):
            table.rows[r].cells[c].text = str(value)

    tbl = table._tbl
    tbl.getparent().remove(tbl)
    current._p.addnext(tbl)

    marker_p = OxmlElement("w:p")
    tbl.addnext(marker_p)
    marker = Paragraph(marker_p, current._parent)
    set_paragraph_text(marker, "")
    return marker


def build_body_blocks() -> list[tuple]:
    return [
        ("h1", "1 引言"),
        ("p", "在现代数据科学领域，机器学习算法已被广泛应用于金融、医疗及房地产等行业的预测分析中。房价预测是一个典型的回归问题，不仅涉及居住面积、建造年份等显性特征，还受到社区环境、房屋质量等级等隐性因素的影响，变量之间往往存在复杂的非线性关系。传统的线性回归方法在高维稀疏数据上表达能力有限，难以充分刻画这种复杂性。"),
        ("p", "集成学习通过组合多个弱学习器来构建强学习器，是处理此类问题的有效途径。Kaggle 竞赛 \"House Prices: Advanced Regression Techniques\" 提供了 79 个原始解释变量和 1460 条训练样本，是验证回归算法的经典基准。本文在该数据集上对比决策树、随机森林与 LightGBM 三种模型：以前两者作为基础算法基线，以 LightGBM 作为扩展算法，系统评估其在特征工程、交叉验证和测试集提交中的表现，为实际房地产估值提供参考。"),
        ("h1", "2 方法或算法"),
        ("p", "本节依次介绍本文实验所涉及的决策树、随机森林、梯度提升决策树（GBDT）及其两个重要扩展实现——XGBoost 与 LightGBM。其中，决策树与随机森林作为基础集成算法用于建立性能基线；GBDT 构成后续扩展算法的理论框架；XGBoost 与 LightGBM 则是在 GBDT 基础上针对训练效率与预测精度所做的改进，本文最终选用 LightGBM 作为主力模型。"),
        ("h2", "2.1 决策树回归"),
        ("p", "决策树（Decision Tree）是一种基于树结构的分而治之（Divide and Conquer）算法。对于回归问题，常用 CART（Classification and Regression Tree）算法构建二叉树：从根节点出发，递归选择最优特征及其分裂点，将样本空间划分为若干互不相交的区域，每个叶节点对应一个预测值，通常取该区域内样本目标值的均值。"),
        ("p", "设数据集包含 N 个样本，当前节点 t 在特征 j 处以阈值 s 分裂为左子节点 tL 和右子节点 tR，回归树的分裂准则一般选取均方误差（MSE）的下降量最大者。节点 t 的均方误差定义为："),
        ("formula", "MSE(t) = (1/|It|) Σ[i∈It](yi - ȳt)²", "(1)"),
        ("p", "其中 It 为落入节点 t 的样本索引集合，ȳt 为该节点样本目标值的均值。最优分裂 (j*, s*) 满足："),
        ("formula", "(j*, s*) = argmin[j,s] [ MSE(tL) + MSE(tR) ]", "(2)"),
        ("p", "决策树的优点在于模型结构直观、可解释性强，且无需对特征做标准化处理。但其主要缺陷在于方差较大：树越深，对训练数据的拟合越充分，也越容易捕获噪声，导致泛化能力下降。因此，单棵决策树通常仅作为更复杂集成方法的基学习器（Base Learner），而非最终预测模型。"),
        ("h2", "2.2 随机森林"),
        ("p", "随机森林（Random Forest, RF）由 Breiman 于 2001 年提出，是一种基于 Bagging（Bootstrap Aggregating）策略的集成学习方法。其核心思想是：对原始训练集进行 B 次有放回抽样（Bootstrap），得到 B 个不同的子训练集；在每个子训练集上独立训练一棵决策树；最终预测时对所有树的输出取平均（回归问题）或投票（分类问题）："),
        ("formula", "ŷ = (1/B) Σ[b=1→B] Tb(x)", "(3)"),
        ("p", "其中 B 为树的数量，Tb(x) 为第 b 棵决策树的预测值。"),
        ("p", "与 Bagging 仅对样本采样不同，随机森林在每次节点分裂时引入特征随机性：从全部 p 个特征中随机抽取 m 个候选特征（通常取 m ≈ √p），仅在这 m 个特征中寻找最优分裂点。这一策略使得各棵树之间的相关性降低，进一步减小集成模型的方差。"),
        ("p", "从偏差-方差分解（Bias-Variance Decomposition）的角度看，随机森林通过集成多棵高方差、低偏差的决策树，以牺牲少量偏差为代价显著降低预测方差，从而提升模型稳定性。此外，随机森林可通过袋外误差（Out-of-Bag Error, OOB Error）估计泛化性能，无需额外划分验证集。在房价预测等特征维度较高、非线性关系显著的任务中，随机森林常作为强有力的基线模型。"),
        ("h2", "2.3 梯度提升决策树（GBDT）"),
        ("p", "梯度提升决策树（Gradient Boosting Decision Tree, GBDT）由 Friedman 于 2001 年提出，属于 Boosting 类集成方法。与随机森林的并行训练不同，GBDT 采用串行方式：每一轮新增一棵回归树，用于拟合当前模型的负梯度（即伪残差），逐步修正前序模型的预测误差。"),
        ("p", "设损失函数为 L(y, F(x))，第 m 轮迭代时，当前模型为 Fm-1(x)，则下一棵回归树 hm(x) 的加入方式可写为："),
        ("formula", "Fm(x) = Fm-1(x) + η · hm(x)", "(4)"),
        ("p", "其中 η ∈ (0, 1] 为学习率（Learning Rate），控制每棵树的贡献权重，较小的学习率配合更多的树通常能获得更好的泛化性能。对于平方损失，有"),
        ("formula", "L = (1/2)(y - F(x))²", ""),
        ("p", "此时负梯度即为普通残差 ri = yi - Fm-1(xi)，因此 hm 实质上是对残差拟合的回归树。"),
        ("p", "GBDT 的核心优势在于：通过加法模型（Additive Model）逐步叠加弱学习器，能够灵活拟合复杂的非线性关系，且对特征缩放不敏感。然而，传统 GBDT 实现在每次分裂时需对所有样本逐特征扫描最优分裂点，当样本量和特征维度较大时，训练效率成为瓶颈。这一局限直接推动了 XGBoost、LightGBM 等高效实现的出现。"),
        ("h2", "2.4 极端梯度提升决策树（XGBoost）"),
        ("p", "极端梯度提升决策树（eXtreme Gradient Boosting Tree, XGBoost）由 Chen 和 Guestrin 于 2016 年提出，是一种基于决策树的集成机器学习算法。XGBoost 使用梯度提升框架，在 GBDT 算法基础上做了多项关键改进，适用于分类和回归问题，在 Kaggle 等数据竞赛中被广泛采用。"),
        ("p", "与原始 GBDT 相比，XGBoost 的核心改进体现在目标函数的设计上。XGBoost 在传统 Boosting 的经验损失基础上，引入了类似 L1/L2 正则化的树复杂度惩罚项。设第 t 轮新增树为 ft，则整体目标函数为："),
        ("formula", "L(t) = Σ[i=1→N] L(yi, ŷi(t-1) + ft(xi)) + Ω(ft)", "(5)"),
        ("p", "其中 Ω(ft) = γT + (1/2)λΣ[j=1→T] wj² 为正则化项，T 为叶节点数，wj 为第 j 个叶节点的权重，γ 和 λ 控制模型复杂度。"),
        ("p", "XGBoost 对损失函数 L 做二阶 Taylor 展开，利用一阶梯度 gi 和二阶 Hessian hi，将目标函数的近似最优解转化为逐特征、逐分裂点的增益（Gain）计算问题，从而高效地搜索最优分裂方向。分裂增益定义为："),
        ("formula", "Gain = (1/2)[ GL²/(HL+λ) + GR²/(HR+λ) - (GL+GR)²/(HL+HR+λ) ] - γ", "(6)"),
        ("p", "其中 GL, HL 分别为左子节点样本的一阶梯度之和与二阶梯度之和，GR, HR 分别为右子节点对应值。"),
        ("p", "此外，XGBoost 还支持列抽样（Column Subsampling），即在每棵树或每次分裂时随机选取部分特征，类似于随机森林的特征随机性，可进一步防止过拟合。上述优点使得 XGBoost 成为各类回归预测任务中较为合适且成熟的手段。本文虽未将 XGBoost 作为最终实验模型，但其在梯度提升框架下的设计思想为理解 LightGBM 提供了重要参照。"),
        ("h2", "2.5 轻量级梯度提升机（LightGBM）"),
        ("p", "轻量级梯度提升机（LightGBM, Light Gradient Boosting Machine）由微软于 2017 年提出，同样基于 GBDT 框架，但在训练效率与内存占用方面做了针对性优化，尤其适用于特征维度高、样本量大的回归与分类任务。LightGBM 可视为与 XGBoost 并列的 GBDT 高效实现，两者在 Kaggle 房价预测等竞赛中均取得了优异表现。"),
        ("p", "LightGBM 的核心改进包括以下三个方面："),
        ("p", "（1）基于直方图的加速算法（Histogram-based Algorithm）。XGBoost 在分裂前需预排序并逐点扫描特征值，而 LightGBM 将连续特征离散化为 k 个桶（bin），在桶边界上寻找最优分裂点。当 k 为常数（如 255）时，分裂搜索的时间复杂度与样本量解耦，训练速度显著提升，内存占用也大幅降低。"),
        ("p", "（2）Leaf-wise 生长策略。XGBoost 默认采用 level-wise（按层）生长，即同一深度的所有叶节点同时分裂；LightGBM 则采用 leaf-wise（按叶）生长，每次选择全局增益最大的叶节点进行分裂。在相同最大深度约束下，leaf-wise 策略通常能以更少的节点达到更低的训练损失，但需通过 max_depth、num_leaves、min_child_samples 等参数控制，以避免过拟合。"),
        ("p", "（3）互斥特征捆绑（Exclusive Feature Bundling, EFB）与单边梯度采样（Gradient-based One-Side Sampling, GOSS）。EFB 将稀疏特征中互斥的列捆绑合并，减少有效特征数；GOSS 保留梯度绝对值较大的样本、对梯度较小的样本随机采样，在几乎不损失精度的前提下减少每轮迭代的数据量。"),
        ("p", "综合上述机制，LightGBM 在保持与 XGBoost 相近预测精度的同时，训练速度可快数倍至数十倍，尤其适合本文所处理的 310 维高维特征场景。本文实验选用的 LightGBM 主要超参数为：学习率 0.03、叶子数 31、最大深度 6、子采样率 0.8、特征采样率 0.8，并在交叉验证中采用早停（Early Stopping, patience=50）策略，在验证集损失不再下降时提前终止训练，以防止过拟合。"),
        ("h1", "3 软件结构和软件实现方法"),
        ("h2", "3.1 开发环境"),
        ("p", "本研究基于 Python 3.10 开发，主要依赖库包括 Pandas（数据处理）、NumPy（数值计算）、Scikit-learn（基线模型与交叉验证）、LightGBM（梯度提升模型）以及 Matplotlib/Seaborn（可视化）。完整依赖见项目根目录 requirements.txt。"),
        ("h2", "3.2 代码结构"),
        ("p", "项目采用模块化设计，核心代码位于 src/ 目录：config.py 管理路径与全局配置；data_process.py 负责数据清洗与特征工程；model_train.py 完成模型训练、十折交叉验证与提交文件生成；visualization.py 生成特征重要性与模型对比图。程序与数据说明详见 README.md。"),
        ("h2", "3.3 实现流程"),
        ("p", "第一步，数据预处理：读取 train.csv 和 test.csv，对缺失值进行语义化填充，构造领域组合特征，对右偏特征和目标变量做对数变换，输出 train_clean.csv（310 维特征）。"),
        ("p", "第二步，模型训练：分别训练决策树、随机森林和 LightGBM，采用十折交叉验证，评价指标为 RMSLE，结果保存至 results/experiment_summary.json。"),
        ("p", "第三步，结果可视化：绘制 LightGBM 特征重要性、三种模型 RMSLE 对比图及预测散点图。"),
        ("p", "第四步，生成报告摘要：自动输出 results/report_summary.md，内容与本文第 5 节实验表格一致。"),
        ("p", "一键运行（推荐）：在项目根目录执行 python run.py。该命令将自动完成上述四个步骤，并校验所有输出文件是否生成。也可使用 run.sh（Linux/macOS）或 run.bat（Windows）。"),
        ("p", "程序代码与数据样例已上传至 GitHub 仓库：https://github.com/BrownSugarMilk/MachineLearning"),
        ("h1", "4 数据描述"),
        ("p", "本文使用的数据集来源于 Kaggle 竞赛 \"House Prices: Advanced Regression Techniques\"。"),
        ("p", "数据规模：训练集 1460 条，测试集 1459 条，原始特征 79 个（不含 Id 和目标变量 SalePrice）。"),
        ("p", "特征类型：包含数值型特征（如 GrLivArea 居住面积、LotArea 地块面积）和分类型特征（如 Neighborhood 社区、HouseStyle 房屋风格）。"),
        ("p", "目标变量：SalePrice（美元），呈明显右偏分布。训练时对目标变量做 log1p 变换，预测结果通过 expm1 还原为原始价格。"),
        ("p", "预处理后：经独热编码与领域特征构造，特征维度扩展至 310 维。新增特征包括 TotalSF（总面积）、TotalBath（总浴室数）、HouseAge（房龄）、RemodAge（改造距今年数）等，有助于模型捕捉面积、房龄与价格之间的非线性关系。"),
        ("h1", "5 实验结果"),
        ("h2", "5.1 实验设置"),
        ("p", "本文采用十折交叉验证（KFold，random_state=42）评估模型性能。由于目标变量经 log1p 变换，对数空间下的 RMSE 等价于 Kaggle 官方评价指标 RMSLE（Root Mean Squared Logarithmic Error）。三种模型的主要参数设置如下："),
        ("bullet", "决策树：max_depth=10"),
        ("bullet", "随机森林：n_estimators=200，max_depth=12，min_samples_leaf=3"),
        ("bullet", "LightGBM：n_estimators=5000，learning_rate=0.03，num_leaves=31，max_depth=6，subsample=0.8，colsample_bytree=0.8，并启用早停（patience=50）"),
        ("h2", "5.2 交叉验证结果"),
        ("p", "表 5-1 汇总了三种模型在十折交叉验证上的平均 RMSLE、标准差、R² 和训练耗时。"),
        ("table", "表 5-1 三种模型十折交叉验证结果对比", ["模型", "平均 RMSLE", "标准差", "R²", "耗时(s)"], [
            ["决策树", "0.2010", "0.0262", "0.7357", "1.43"],
            ["随机森林", "0.1398", "0.0234", "0.8724", "7.38"],
            ["LightGBM", "0.1243", "0.0203", "0.8993", "18.29"],
        ]),
        ("p", "LightGBM 十折交叉验证各折 RMSLE 如表 5-2 所示。"),
        ("table", "表 5-2 LightGBM 十折交叉验证各折 RMSLE", ["折次", "1", "2", "3", "4", "5", "6", "7", "8", "9", "10"], [
            ["RMSLE", "0.1217", "0.1455", "0.1025", "0.1205", "0.1488", "0.1391", "0.1402", "0.1114", "0.1319", "0.0811"],
        ]),
        ("h2", "5.3 测试集提交"),
        ("p", "将 LightGBM 模型在全部训练集上训练后，对测试集进行预测并提交至 Kaggle 平台。提交文件格式为 Id, SalePrice，程序自动生成于 results/submission_lightgbm.csv。Kaggle 官方采用 RMSLE 评价测试集表现，该指标与本文交叉验证使用的指标一致，便于横向对比。"),
        ("h2", "5.4 特征重要性分析"),
        ("p", "LightGBM 特征重要性排名前五位为：TotalSF（总面积）、LotArea（地块面积）、OverallCond（整体状况）、GrLivArea（地上居住面积）和 BsmtFinSF1（地下室装修面积）。面积类与房屋状况特征排名靠前，与房地产领域常识一致，表明模型学到了合理的规律。"),
        ("h1", "6 实验分析与总结"),
        ("h2", "6.1 结果分析"),
        ("p", "从表 5-1 可以看出，三种模型的性能呈现明显的递进关系：决策树 RMSLE 为 0.2010，随机森林降至 0.1398，LightGBM 进一步降至 0.1243。决策树作为单模型，方差较大且容易在训练集上过拟合，R² 仅为 0.7357；随机森林通过集成多棵树有效降低了方差，R² 提升至 0.8724；LightGBM 采用 Boosting 策略和 leaf-wise 生长，R² 达到 0.8993，在相同特征工程条件下取得了最优的交叉验证表现。"),
        ("p", "LightGBM 相比随机森林 RMSLE 相对降低约 11.1%（从 0.1398 到 0.1243），说明梯度提升方法在该数据集上能更精细地拟合特征与价格之间的非线性关系。LightGBM 训练耗时（18.29s）高于随机森林（7.38s），但考虑到其性能提升和早停机制带来的效率优化，这一代价是合理的。"),
        ("h2", "6.2 不足与改进方向"),
        ("p", "（1）特征工程仍有提升空间，可进一步构造多项式交互特征或采用目标编码（Target Encoding）处理高基数类别变量。"),
        ("p", "（2）超参数可通过 Grid Search 或贝叶斯优化系统搜索，当前参数基于经验设置。"),
        ("p", "（3）可尝试 Stacking 集成，将 LightGBM 与随机森林的预测结果融合，有望进一步降低 RMSLE。"),
        ("h2", "6.3 总结"),
        ("p", "本文基于 Kaggle 房价预测竞赛数据集，以决策树和随机森林为基线、LightGBM 为扩展算法，完成了从数据预处理、特征工程、模型训练到结果可视化的完整流程。实验表明，经特征优化与参数调优后的 LightGBM 在十折交叉验证中取得了 0.1243 的 RMSLE，优于两种基础集成方法，验证了扩展算法在复杂回归问题中的有效性与实用价值。"),
        ("h1", "参考文献"),
        ("ref", "[1] Ke G, Meng Q, Finley T, et al. LightGBM: A Highly Efficient Gradient Boosting Decision Tree[C]. Advances in Neural Information Processing Systems, 2017: 3146-3154."),
        ("ref", "[2] Chen T, Guestrin C. XGBoost: A Scalable Tree Boosting System[C]. Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining, 2016: 785-794."),
        ("ref", "[3] Breiman L. Random Forests[J]. Machine Learning, 2001, 45(1): 5-32."),
        ("ref", "[4] Friedman J H. Greedy Function Approximation: A Gradient Boosting Machine[J]. Annals of Statistics, 2001, 29(5): 1189-1232."),
        ("ref", "[5] Pedregosa F, et al. Scikit-learn: Machine Learning in Python[J]. Journal of Machine Learning Research, 2011, 12: 2825-2830."),
    ]


def insert_body_content(doc: Document, anchor: Paragraph):
    current = anchor
    for block in build_body_blocks():
        kind = block[0]
        if kind == "h1":
            current = insert_paragraph_after(current, doc)
            set_paragraph_text(current, block[1], style="Heading 1", bold=True)
        elif kind == "h2":
            current = insert_paragraph_after(current, doc)
            set_paragraph_text(current, block[1], bold=True)
        elif kind == "p":
            current = insert_paragraph_after(current, doc)
            set_paragraph_text(current, block[1])
        elif kind == "bullet":
            current = insert_paragraph_after(current, doc)
            set_paragraph_text(current, block[1], style="List Paragraph")
        elif kind == "formula":
            current = insert_paragraph_after(current, doc)
            text = f"{block[1]}    {block[2]}".strip()
            set_paragraph_text(current, text, align_center=True)
        elif kind == "ref":
            current = insert_paragraph_after(current, doc)
            set_paragraph_text(current, block[1])
        elif kind == "table":
            _, caption, headers, rows = block
            current = insert_table_after(doc, current, headers, rows, caption)


def fill_report_docx(template_path: Path = TEMPLATE_PATH, output_path: Path = OUTPUT_PATH):
    if not template_path.exists():
        if BACKUP_SOURCE.exists():
            shutil.copy2(BACKUP_SOURCE, template_path)
        else:
            raise FileNotFoundError(f"未找到 Word 模板：{template_path}")

    doc = Document(str(template_path))
    fill_cover_page(doc)
    anchor = clear_body_after_intro(doc)
    insert_body_content(doc, anchor)
    doc.save(str(output_path))
    return output_path


def main():
    output = fill_report_docx()
    print(f"Word 报告已生成：{output}")


if __name__ == "__main__":
    main()
